import os
import sys
import subprocess

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("python-docx is not installed. Installing it now...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
        print("python-docx successfully installed!")
    except Exception as e:
        print(f"Failed to install python-docx: {e}")
        sys.exit(1)

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    kwargs: top, bottom, left, right, insideH, insideV
    value is a dictionary: {'sz': 12, 'val': 'single', 'color': 'FF0000', 'space': '0'}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def create_letter():
    doc = Document()
    
    # 1. Page Margins Setup (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # 2. Base Typography Setup (Times New Roman, 12pt)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Helper to add standard paragraph
    def add_para(text="", space_after=12, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        if text:
            p.add_run(text)
        return p

    # 3. Header Letterhead Box (Dotted Border Box)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    # Apply light gray dotted border to the box
    set_cell_border(
        cell,
        top={'sz': 6, 'val': 'dotted', 'color': '888888', 'space': '0'},
        bottom={'sz': 6, 'val': 'dotted', 'color': '888888', 'space': '0'},
        left={'sz': 6, 'val': 'dotted', 'color': '888888', 'space': '0'},
        right={'sz': 6, 'val': 'dotted', 'color': '888888', 'space': '0'}
    )
    
    # Add text inside the box
    p_box = cell.paragraphs[0]
    p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_box.paragraph_format.space_before = Pt(8)
    p_box.paragraph_format.space_after = Pt(8)
    r_box = p_box.add_run("[Law Firm / Individual Lawyer Letterhead]\n(If printing on official lawyer letterhead, discard this box)")
    r_box.italic = True
    r_box.font.size = Pt(10)
    r_box.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
    
    # Spacer below box
    add_para(space_after=24)
    
    # 4. Document Title
    p_title = add_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    r_title = p_title.add_run("LETTER OF SUPPORT")
    r_title.bold = True
    r_title.font.size = Pt(14)
    
    # Subtitle
    p_sub = add_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    r_sub = p_sub.add_run("(For University Defense)")
    r_sub.italic = True
    r_sub.font.size = Pt(11)
    
    # 5. Date
    p_date = add_para(space_after=18)
    r_date_label = p_date.add_run("Date: ")
    r_date_label.bold = True
    p_date.add_run("____________________")
    
    # 6. Recipient
    add_para("To Whom It May Concern,", space_after=18)
    
    # 7. Subject
    p_subj = add_para(space_after=18)
    r_subj_label = p_subj.add_run("Subject: ")
    r_subj_label.bold = True
    r_subj_val = p_subj.add_run("Letter of Support for Barq-e-Insaf Project")
    r_subj_val.bold = True
    
    # 8. Body Paragraph 1
    p_body1 = add_para(space_after=12)
    p_body1.add_run("This letter is to confirm that [Law Firm Name / Lawyer Name] is aware of and supports the ")
    r_proj = p_body1.add_run("Barq-e-Insaf")
    r_proj.bold = True
    p_body1.add_run(" project being developed by ")
    r_names = p_body1.add_run("Anjum Kalwar, Momina Amin, and Rezan Sohail")
    r_names.bold = True
    p_body1.add_run(" as their Final Year Project at Bahria University, Karachi Campus.")
    
    # Body Paragraph 2
    p_body2 = add_para(space_after=12)
    p_body2.add_run("We understand that ")
    p_body2.add_run("Barq-e-Insaf").bold = True
    p_body2.add_run(
        " is an AI-powered mobile application designed to improve access to justice in Sindh by providing "
        "legal guidance through a multilingual chatbot, connecting citizens with verified lawyers, and "
        "enabling secure digital evidence management for property and family law disputes."
    )
    
    # Body Paragraph 3
    p_body3 = add_para(space_after=12)
    p_body3.add_run("We have reviewed the project concept and believe it addresses a critical need in our legal system. We are willing to:")
    
    # 9. Bullet Points
    bullets = [
        "Provide guidance on legal procedures and documentation related to property and family laws in Sindh.",
        "Review and validate the legal accuracy of the chatbot's responses and legal content.",
        "Act as a reference/advisor for the project team during development.",
        "Potentially use the platform upon completion to connect with clients."
    ]
    
    for b in bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        p_b.paragraph_format.line_spacing = 1.15
        p_b.paragraph_format.left_indent = Inches(0.5)
        p_b.add_run(b)
        
    # Spacer
    add_para(space_after=12)
    
    # Body Paragraph 4
    add_para("We support this initiative and encourage the team to continue their work on this impactful project.", space_after=24)
    
    # Closing
    add_para("Sincerely,", space_after=36)
    
    # 10. Lawyer Signature Form Fields
    p_sig = add_para(space_after=8)
    p_sig.add_run("Signature: ").bold = True
    p_sig.add_run("_____________________________")
    
    p_name = add_para(space_after=8)
    p_name.add_run("Name: ").bold = True
    p_name.add_run("_________________________________________")
    
    p_desig = add_para(space_after=8)
    p_desig.add_run("Designation: ").bold = True
    p_desig.add_run("_________________________________________")
    
    p_firm = add_para(space_after=8)
    p_firm.add_run("Law Firm Name: ").bold = True
    p_firm.add_run("________________________________________")
    
    p_reg = add_para(space_after=8)
    p_reg.add_run("Sindh Bar Council Registration No.: ").bold = True
    p_reg.add_run("____________________")
    
    p_phone = add_para(space_after=8)
    p_phone.add_run("Contact Number: ").bold = True
    p_phone.add_run("____________________________________")
    
    p_email = add_para(space_after=18)
    p_email.add_run("Email: ").bold = True
    p_email.add_run("___________________________________________")
    
    # Official Stamp Box placeholder
    p_stamp = add_para(space_after=12)
    r_stamp = p_stamp.add_run("[ Official Stamp ]")
    r_stamp.italic = True
    r_stamp.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
    
    # Save the file in the requested project directory
    target_dir = r"C:\Users\HP\OneDrive\Documents\GitHub\barq-e-insaf"
    output_path = os.path.join(target_dir, "Letter_of_Support.docx")
    
    try:
        doc.save(output_path)
        print(f"SUCCESS: Document saved as {output_path}")
    except Exception as e:
        # Fallback to local directory if permissions fail
        local_path = "Letter_of_Support.docx"
        doc.save(local_path)
        print(f"FALLBACK: Saved to {os.path.abspath(local_path)} due to: {e}")

if __name__ == "__main__":
    create_letter()
